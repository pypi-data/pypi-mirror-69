import re
from docx import Document

class get_log:
    def __init__(self,files,month_list):
        self.files=files
        self.month_list=month_list

    def get_log(self):
        list_hostname = []
        list_log = []
        for file in self.files:
            try:
                print('Processing File :')
                print(file)
                read_file = open(file, 'r')
                read_file_list = read_file.readlines()

                for file in read_file_list:
                    if re.findall('^hostname (.*)', file):
                        hostname = re.findall('^hostname (.*)', file)

                    for month in self.month_list:
                        if re.findall('^.*'+month+'.*-[0,1,2,3,4]-.*', file):  
                            log = re.findall('^.*'+month+'.*-[0,1,2,3,4]-.*', file)
                            list_hostname.append(hostname)
                            list_log.append(log)
            #except NameError:
            # raise
            except:
                pass
        '''count_host = 0
        for hostname in list_hostname:
            print(hostname+' '+list_log[count_host])
            count_host+=1'''

        print(list_hostname)
        print(list_log)

        #using document docx module
        document = Document()

        #add to document
        p = document.add_paragraph('')
        p.add_run('LOG SUMMARY').bold = True
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Log Event'
        hdr_cells[2].text = 'Severity'
        count_host = 0        
        for hostname in list_hostname:
            row_cells = table.add_row().cells
            row_cells[0].text = hostname
            row_cells[1].text = list_log[count_host]
            count_host+=1

        #save document
        print('Saving Document')
        document.save('show_log.docx')
        print('Document has been saved to show_log.docx')