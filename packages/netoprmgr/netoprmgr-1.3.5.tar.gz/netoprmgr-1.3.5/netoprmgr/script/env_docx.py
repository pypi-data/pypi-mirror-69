from docx import Document
import sqlite3

class env_docx:
    @staticmethod
    def env_docx():
        #open db connection
        db = sqlite3.connect('env_pmdb')
        cursor = db.cursor()
        #using document docx module
        document = Document()
        
        #SHOW ENVIRONMENT
        #sql query
        cursor.execute('''SELECT * FROM envtable''')
        records = cursor.fetchall()
        #add to document
        p = document.add_paragraph('')
        p.add_run('Hardware Condition Analysis').bold = True
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'System'
        hdr_cells[2].text = 'Item'   
        hdr_cells[3].text = 'Status'       
        for row in records:
            row_cells = table.add_row().cells
            row_cells[0].text = (row[1])
            row_cells[1].text = (row[2]) 
            row_cells[2].text = (row[3])
            row_cells[3].text = (row[4])
        
        #close database
        db.close()
        
        #save document
        print('Saving Document')
        document.save('show_environment.docx')
        print('Document has been saved to show_environment.docx')