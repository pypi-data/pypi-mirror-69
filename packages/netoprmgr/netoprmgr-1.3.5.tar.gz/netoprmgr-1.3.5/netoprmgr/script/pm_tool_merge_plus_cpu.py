import os
from docx import Document
import re
import traceback

#file identification method
def file_identification():
    regex_file = re.findall(".py$", file)
    if regex_file:
        pass
    else:
        try:
            #debug file that currently opened
            print('file opened right now -----------------------------------------------')
            print(file)
            read_file=open(file,'r', encoding='utf8')
            read_file_list=read_file.readlines()
            for every_line in read_file_list:
                #identification for IOS XE
                if 'Cisco IOS XE Software' in every_line:                  
                    #SOFTWARE
                    #get device name method
                    get_device_name()
                    #get model using get_data_type_one method
                    get_data_type_one(every_words[word_adjust],'bytes of memory',0,'cisco','Cisco',1)
                    #append data to list
                    list_model.append(every_words[word_adjust])
                    #get ios version using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'System image file is',0,'System',4)
                    #append data to list
                    list_ios_version.append(every_words[word_adjust])
                    #get uptime using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'hours',0,'up',1)
                    #append data to list
                    list_uptime.append(every_words[word_adjust+2]+' '+every_words[word_adjust+3]+' '+every_words[word_adjust+4]+' '+every_words[word_adjust+5])
                    #get configuration register(conf reg) using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'Configuration register',0,'Configuration',3)
                    #append data to list
                    list_config_register.append(every_words[word_adjust])                    
                    #get version using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'Cisco IOS Software',0,'Version',1)
                    #append data to list
                    list_version.append(every_words[word_adjust+count_word_after])
                    #CPU
                    #get cpu using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'CPU utilization',0,'seconds:',1)
                    #append data to list
                    regex_word=every_words[word_adjust+count_word_after]
                    regex_word_cut = re.findall("(.*)/.*$", regex_word)
                    list_cpu.append(regex_word_cut)
                    #get cpu using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'CPU utilization',0,'seconds:',1)
                    #append data to list
                    regex_word=every_words[word_adjust+count_word_after]
                    regex_word_cut_two = re.findall(".*/(.*);$", regex_word)
                    list_cpu_interrupt.append(regex_word_cut_two)
                    #regexing pusing
                    str_cpu = str(regex_word_cut)
                    str_cpu_interrupt = str(regex_word_cut_two)
                    regex_cpu = re.findall("^.*'(.*)%']$", str_cpu)
                    regex_cpu_interrupt = re.findall("^.*'(.*)%']$", str_cpu_interrupt)
                    int_cpu = [int(x) for x in regex_cpu]
                    int_cpu_interrupt = [int(x) for x in regex_cpu_interrupt]
                    total_cpu=int_cpu[0]-int_cpu_interrupt[0]
                    list_cpu_total.append(str(total_cpu)+'%')
                    #if condition for status
                    if total_cpu<21 :
                        status='Low'
                    elif total_cpu<81 :
                        status='Medium'
                    else:
                        status='High'
                    #append data to list    
                    list_cpu_status.append(status)                     
                    #MEMORY
                    #get memory using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'Processor Pool Total:',0,'Total:',1)
                    find_word=every_words[word_adjust+count_word_after]
                    find_word_two=every_words[word_adjust+count_word_after+2]
                    memory_percentage=(int(find_word_two)/int(find_word))*100
                    print('memory_percentage')
                    print(memory_percentage)
                    regex_memory = re.findall(".*", str(memory_percentage))
                    print('regex_memory')
                    print(regex_memory)
                    #append data to list
                    list_memory.append(str(memory_percentage)+'%')
                    #get top 3 memory using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'show process memory sorted',6,'Total:',0)
                    find_word=every_words[word_adjust+7]
                    get_data_type_two(every_words[word_adjust],'show process memory sorted',7,'Total:',0)
                    find_word_two=every_words[word_adjust+7]
                    get_data_type_two(every_words[word_adjust],'show process memory sorted',8,'Total:',0)
                    find_word_three=every_words[word_adjust+7]
                    memory_top_three=(find_word+'\n'+find_word_two+'\n'+find_word_three)
                    print('memory_top_three')
                    print(memory_top_three)
                    #append data to list
                    list_memory_top_three.append(memory_top_three)
                    #if condition for status
                    if memory_percentage<21 :
                        status='Low'
                    elif memory_percentage<81 :
                        status='Medium'
                    else:
                        status='High'
                    print('status')
                    print(status)
                    list_memory_status.append(status)
                    break
                #identification for IOS
                elif 'Cisco IOS Software' in every_line:
                    #SOFTWARE
                    #get device name method
                    get_device_name()
                    #get model using get_data_type_one method
                    get_data_type_one(every_words[word_adjust],'If you require further assistance please contact us by sending email',3,'cisco','Cisco',1)
                    #append data to list
                    list_model.append(every_words[word_adjust])
                    #get ios version using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'System image file is',0,'System',4)
                    #append data to list
                    list_ios_version.append(every_words[word_adjust])
                    #get uptime using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'uptime',0,'uptime',2)
                    #append data to list
                    list_uptime.append(every_words[word_adjust+1]+' '+every_words[word_adjust+2]+' '+every_words[word_adjust+3]+' '+every_words[word_adjust+4])
                    #get configuration register(conf reg) using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'Configuration register',0,'Configuration',3)
                    #append data to list
                    list_config_register.append(every_words[word_adjust])
                    #get version using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'Cisco IOS Software',0,'Version',1)
                    #append data to list
                    list_version.append(every_words[word_adjust+count_word_after])
                    #CPU
                    #get cpu using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'CPU utilization',0,'seconds:',1)
                    #append data to list
                    regex_word=every_words[word_adjust+count_word_after]
                    regex_word_cut = re.findall("(.*)/.*$", regex_word)
                    list_cpu.append(regex_word_cut)                    
                    #get cpu using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'CPU utilization',0,'seconds:',1)
                    #append data to list
                    regex_word=every_words[word_adjust+count_word_after]
                    regex_word_cut_two = re.findall(".*/(.*);$", regex_word)
                    list_cpu_interrupt.append(regex_word_cut_two)                    
                    #regexing pusing
                    str_cpu = str(regex_word_cut)
                    str_cpu_interrupt = str(regex_word_cut_two)
                    regex_cpu = re.findall("^.*'(.*)%']$", str_cpu)
                    regex_cpu_interrupt = re.findall("^.*'(.*)%']$", str_cpu_interrupt)
                    int_cpu = [int(x) for x in regex_cpu]
                    int_cpu_interrupt = [int(x) for x in regex_cpu_interrupt]
                    total_cpu=int_cpu[0]-int_cpu_interrupt[0]
                    list_cpu_total.append(str(total_cpu)+'%')                    
                    #if condition for status
                    if total_cpu<21 :
                        status='Low'
                    elif total_cpu<81 :
                        status='Medium'
                    else:
                        status='High'
                    #appen data to list
                    list_cpu_status.append(status)
                    #MEMORY
                    #get memory using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'Processor Pool Total:',0,'Total:',1)
                    find_word=every_words[word_adjust+count_word_after]
                    find_word_two=every_words[word_adjust+count_word_after+2]
                    memory_percentage=(int(find_word_two)/int(find_word))*100
                    print('memory_percentage')
                    print(memory_percentage)
                    regex_memory = re.findall(".*", str(memory_percentage))
                    print('regex_memory')
                    print(regex_memory)
                    #append data to list
                    list_memory.append(str(memory_percentage)+'%')
                    #get top 3 memory using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'show process memory sorted',6,'Total:',0)
                    find_word=every_words[word_adjust+7]
                    get_data_type_two(every_words[word_adjust],'show process memory sorted',7,'Total:',0)
                    find_word_two=every_words[word_adjust+7]
                    get_data_type_two(every_words[word_adjust],'show process memory sorted',8,'Total:',0)
                    find_word_three=every_words[word_adjust+7]
                    memory_top_three=(find_word+'\n'+find_word_two+'\n'+find_word_three)
                    print('memory_top_three')
                    print(memory_top_three)
                    #append data to list
                    list_memory_top_three.append(memory_top_three)
                    #if condition for status
                    if memory_percentage<21 :
                        status='Low'
                    elif memory_percentage<81 :
                        status='Medium'
                    else:
                        status='High'
                    print('status')
                    print(status)
                    list_memory_status.append(status)                    
                    break
                #identification for ASA
                elif 'Cisco Adaptive Security Appliance Software' in every_line:
                    #get device name method
                    get_device_name()
                    #get model using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'Hardware:',0,'Hardware:',1)
                    #append data to list
                    list_model.append(every_words[word_adjust])
                    #get ios version using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'System image file is',0,'System',4)
                    #append data to list
                    list_ios_version.append(every_words[word_adjust])
                    #get uptime using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'hours',0,'up',1)
                    #append data to list
                    list_uptime.append(every_words[word_adjust+1]+' '+every_words[word_adjust+2]+' '+every_words[word_adjust+3]+' '+every_words[word_adjust+4])
                    #get configuration register(conf reg) using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'Configuration register',0,'Configuration',3)
                    #append data to list
                    list_config_register.append(every_words[word_adjust])
                    #get version using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'Cisco Adaptive Security Appliance Software',0,'Version',1)
                    #append data to list
                    list_version.append(every_words[word_adjust+count_word_after])
                    #CPU
                    #get cpu using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'CPU utilization',0,'seconds:',1)
                    #append data to list
                    regex_word=every_words[word_adjust+count_word_after]
                    regex_word_cut = re.findall("(.*);$", regex_word)
                    list_cpu.append(regex_word_cut)                    
                    #append data to list, without any regex
                    list_cpu_interrupt.append(['0%'])                    
                    #regexing pusing
                    str_cpu = str(regex_word_cut)
                    regex_cpu = re.findall("^.*'(.*)%']$", str_cpu)
                    int_cpu = [int(x) for x in regex_cpu]
                    total_cpu=int_cpu[0]                    
                    #append data to list
                    list_cpu_total.append(str(total_cpu)+'%')                    
                    #if condition for status
                    if total_cpu<21 :
                        status='Low'
                    elif total_cpu<81 :
                        status='Medium'
                    else:
                        status='High'
                    #append data to list
                    list_cpu_status.append(status)
                    
                    #MEMORY
                    #get memory using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'Used memory:',0,'bytes',1)
                    find_word=every_words[word_adjust+count_word_after]
                    str_find_word=str(find_word)
                    regex_word = re.findall("[0-9][0-9]", str_find_word)
                    int_word = [int(x) for x in regex_word]
                    status_word=int_word[0]
                    print('find_word_ASAAAA')
                    print(int_word)
                    print(regex_word)
                    #append data to list
                    list_memory.append(str(int_word)+'%')
                    #get top 3 memory using get_data_type_one method
                    #append data to list
                    list_memory_top_three.append('-')
                    #if condition for status
                    if status_word<21 :
                        status='Low'
                    elif status_word<81 :
                        status='Medium'
                    else:
                        status='High'
                    print('status')
                    print(status)
                    list_memory_status.append(status)                    
                    break
                #identification for Nexus
                elif 'Cisco Nexus Operating System (NX-OS)' in every_line:
                    #get device name method
                    get_device_name()
                    #get model using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'Hardware',1,'cisco',0)
                    #append data to list
                    list_model.append(every_words[word_adjust+1]+' '+every_words[word_adjust+2])
                    #get ios version using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'system image file is:',0,'system',4)
                    #append data to list
                    list_ios_version.append(every_words[word_adjust])
                    #get uptime using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'uptime',0,'uptime',1)
                    #append data to list
                    list_uptime.append(every_words[word_adjust+2]+' '+every_words[word_adjust+3]+' '+every_words[word_adjust+4]+' '+every_words[word_adjust+5])
                    #append data to list
                    list_config_register.append('-')
                    #get version using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'system:',0,'system:',2)
                    #append data to list
                    list_version.append(every_words[word_adjust+count_word_after])
                    #CPU
                    #get cpu using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'CPU util',0,':',1)
                    #regex
                    regex_word=every_words[word_adjust+count_word_after]
                    regex_word_two=every_words[word_adjust+count_word_after+2]
                    str_cpu_user = str(regex_word)
                    str_cpu_kernel = str(regex_word_two)
                    regex_cpu_user = re.findall("^(.*)%$", str_cpu_user)
                    regex_cpu_kernel = re.findall("^(.*)%$", str_cpu_kernel)
                    float_cpu_user = [float(x) for x in regex_cpu_user]
                    float_cpu_kernel = [float(x) for x in regex_cpu_kernel]
                    total_cpu=float_cpu_user[0]+float_cpu_kernel[0]
                    #append data to list
                    list_cpu.append(str(total_cpu)+'%')
                    list_cpu_interrupt.append('0%')
                    list_cpu_total.append(str(total_cpu)+'%')                   
                    #if condition for status
                    if total_cpu<21 :
                        status='Low'
                    elif total_cpu<81 :
                        status='Medium'
                    else:
                        status='High'
                    #append data to list
                    list_cpu_status.append(status)
                    #MEMORY
                    #get memory using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'Memory usage:',0,'Total:',0)
                    find_word=every_words[word_adjust+count_word_after]
                    get_data_type_two(every_words[word_adjust],'Memory usage:',0,'Total:',2)
                    find_word_two=every_words[word_adjust+count_word_after]
                    regex_word=re.sub("\D", "", find_word)
                    regex_word_two=re.sub("\D", "", find_word_two)
                    int_word=int(regex_word)
                    int_word_two=int(regex_word_two)
                    memory_percentage=(int_word_two/int_word)*100
                    #append datat to list
                    list_memory.append(str(memory_percentage)+'%')
                    #append data to list
                    list_memory_top_three.append('-')
                    #if condition for status
                    if memory_percentage<21 :
                        status='Low'
                    elif memory_percentage<81 :
                        status='Medium'
                    else:
                        status='High'
                    print('status')
                    print(status)
                    list_memory_status.append(status)
                    
                    break
                #identification for WLC
                elif 'Series Wireless LAN Controller' in every_line:
                    #get device name method
                    get_device_name()
                    #get model using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'NAME: "Chassis"',0,'NAME',4)
                    #append data to list
                    list_model.append(every_words[word_adjust]+' '+every_words[word_adjust+1]+' '+every_words[word_adjust+2]+' '+every_words[word_adjust+3])
                    #get ios version using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'Product Version',0,'Product Version',2)
                    #append data to list
                    list_ios_version.append(every_words[word_adjust])
                    #get uptime using get_data_type_two method
                    get_data_type_two(every_words[word_adjust],'System Up Time',0,'System',2)
                    #append data to list
                    list_uptime.append(every_words[word_adjust+1]+' '+every_words[word_adjust+2]+' '+every_words[word_adjust+3]+' '+every_words[word_adjust+4])
                    #append data to list
                    list_config_register.append('-')
                    #get version using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'Product Version',0,'Product Version',2)
                    #append data to list
                    list_version.append(every_words[word_adjust])
                    #CPU
                    #get cpu using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'CPU Current Usage',0,'Usage............................. ',3)
                    regex_word=every_words[word_adjust+count_word_after]
                    total_cpu=int(regex_word)
                    #append data to list
                    list_cpu.append(str(regex_word)+'%')
                    list_cpu_interrupt.append('0%')
                    list_cpu_total.append(str(regex_word)+'%')                    
                    #if condition for status
                    if total_cpu<21 :
                        status='Low'
                    elif total_cpu<81 :
                        status='Medium'
                    else:
                        status='High'
                    #append data to list    
                    list_cpu_status.append(status)
                    #MEMORY
                    #get memory using get_data_type_one method
                    get_data_type_two(every_words[word_adjust],'Memory Average Usage',0,'Usage............................. ',3)
                    find_word=every_words[word_adjust+count_word_after]
                    int_word=int(find_word)
                    print('int_word')
                    print(int_word)
                    #append data to list
                    list_memory.append(str(int_word)+'%')
                    #get top 3 memory using get_data_type_one method
                    #append data to list
                    list_memory_top_three.append('-')
                    #if condition for status
                    if int_word<21 :
                        status='Low'
                    elif int_word<81 :
                        status='Medium'
                    else:
                        status='High'
                    print('status')
                    print(status)
                    list_memory_status.append(status)
                    break
        #except:
            #pass
        except NameError:
            raise

#get data methods
def get_device_name():
    global list_device_name
    regex_word = re.findall("(.*).txt.*$", file)    
    list_device_name.append(regex_word)

def get_data_type_one(model,word_in_file,row_adjust,word_in_line,word_in_line2,word_adjustment):
    #global list_model
    global every_words
    global count_word
    global word_adjust
    global count_word_after
    read_file=open(file,'r', encoding='utf8')
    read_file_list=read_file.readlines()
    count_line=0
    for every_line in read_file_list:
        word_in_file = word_in_file #isi
        if word_in_file in every_line:
            count_word=0
            row_adjust=row_adjust #isi
            next_line=read_file_list[count_line+row_adjust]
            every_words=next_line.split()
            for every_word in every_words:
                print('word parsing')
                print(every_word)
                word_in_line = word_in_line #isi
                word_in_line2 = word_in_line2 #isi
                word_adjust = word_adjustment #isi
                if word_in_line in every_word:
                    model=model
                    count_word_after=count_word
                elif word_in_line2 in every_word:
                    model=model
                    count_word_after=count_word
                count_word+=1
            break   
        count_line+=1   
    read_file.close

def get_data_type_two(model,word_in_file,row_adjust,word_in_line,word_adjustment):
    #global list_model
    global every_words
    global count_word
    global word_adjust
    global count_word_after
    read_file=open(file,'r', encoding='utf8')
    read_file_list=read_file.readlines()
    count_line=0
    for every_line in read_file_list:
        word_in_file = word_in_file #isi
        if word_in_file in every_line:
            count_word=0
            row_adjust=row_adjust #isi
            next_line=read_file_list[count_line+row_adjust]
            every_words=next_line.split()
            for every_word in every_words:
                print('word parsing')
                print(every_word)
                word_in_line = word_in_line #isi
                word_adjust = word_adjustment #isi
                if word_in_line in every_word:                    
                    model=model
                    count_word_after=count_word
                count_word+=1
            break            
        count_line+=1        
    read_file.close

#using document docx module
document = Document()
#get current directory file
print('test--------------------------------------------------------------')
current_directory = os.path.dirname(os.path.realpath(__file__))
files = os.listdir(current_directory)
#error log for further analisys
file_name = 'error.log'
write_error_output = open(file_name,'w')
#all appended list
list_device_name=[]
list_model=[]
list_ios_version=[]
list_uptime=[]
list_config_register=[]
list_version=[]
list_cpu=[]
list_cpu_interrupt=[]
list_cpu_total=[]
list_cpu_status=[]
list_memory=[]
list_memory_top_three=[]
list_memory_status=[]
#global variable
every_words=['first list']
count_word=0
count_word_after=0
word_adjust=0
#debug all files available
print('all files')
print(files)
#loop over files
for file in files:
    file_identification()
#new variable
list_total=len(list_version)
list_arrange=set(list_version)
#print for checking
print('list_memory_status')
print(list_memory_status)
#adding list to docx table
#software table summary
p = document.add_paragraph('')
p.add_run('SOFTWARE TABLE SUMMARY').bold = True

table = document.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Version'
hdr_cells[1].text = 'Total'
hdr_cells[2].text = 'Percentage'
count_cell=0
for a in list_arrange:
    var_total=list_version.count(a)
    row_cells = table.add_row().cells
    row_cells[0].text = a
    row_cells[1].text = str(var_total) 
    row_cells[2].text = str((var_total/list_total)*100)
    count_cell+=1
#software table
p = document.add_paragraph('')
p.add_run('SOFTWARE TABLE').bold = True

table = document.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Device Name'
hdr_cells[1].text = 'Model'
hdr_cells[2].text = 'IOS Version'
hdr_cells[3].text = 'Uptime'
hdr_cells[4].text = 'Config Register'
count_cell=0
for a in list_device_name:

    row_cells = table.add_row().cells
    row_cells[0].text = list_device_name[count_cell]
    row_cells[1].text = list_model[count_cell] 
    row_cells[2].text = list_ios_version[count_cell]
    row_cells[3].text = list_uptime[count_cell]
    row_cells[4].text = list_config_register[count_cell]
    count_cell+=1
#cpu summary table
p = document.add_paragraph('')
p.add_run('CPU SUMMARY TABLE').bold = True

table = document.add_table(rows=1, cols=6)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Device Name'
hdr_cells[1].text = 'Model'
hdr_cells[2].text = 'Total'
hdr_cells[3].text = 'Process'
hdr_cells[4].text = 'Interrupt'
hdr_cells[5].text = 'Status'
count_cell=0
for a in list_device_name:

    row_cells = table.add_row().cells
    row_cells[0].text = list_device_name[count_cell]
    row_cells[1].text = list_model[count_cell] 
    row_cells[2].text = list_cpu[count_cell]
    row_cells[3].text = list_cpu_total[count_cell]
    row_cells[4].text = list_cpu_interrupt[count_cell]
    row_cells[5].text = list_cpu_status[count_cell]
    count_cell+=1
    
#memory summary table
p = document.add_paragraph('')
p.add_run('MEMORY SUMMARY TABLE').bold = True

table = document.add_table(rows=1, cols=5)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Device Name'
hdr_cells[1].text = 'Model'
hdr_cells[2].text = 'Processor Memory Utilization'
hdr_cells[3].text = 'Top Process'
hdr_cells[4].text = 'Status'
count_cell=0
for a in list_device_name:

    row_cells = table.add_row().cells
    row_cells[0].text = list_device_name[count_cell]
    row_cells[1].text = list_model[count_cell] 
    row_cells[2].text = list_memory[count_cell]
    row_cells[3].text = list_memory_top_three[count_cell]
    row_cells[4].text = list_memory_status[count_cell]
    count_cell+=1
    
#save document
document.save('Preventive Maintenance.docx')
print('Document has been saved to Preventive Maintenance.docx')