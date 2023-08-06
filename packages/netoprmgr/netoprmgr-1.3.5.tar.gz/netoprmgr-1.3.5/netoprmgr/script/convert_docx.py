from docx import Document
import sqlite3
import re

class convert_docx:
    @staticmethod
    def convert_docx():
        #open db connection
        db = sqlite3.connect('pmdb')
        cursor = db.cursor()
        #using document docx module
        document = Document()
        
        #SOFTWARE SUMMARY TABLE
        #sql query
        cursor.execute('''SELECT version, COUNT(*) FROM swsumtable GROUP BY version''')
        records = cursor.fetchall()
        cursor.execute('''SELECT COUNT(version) FROM swsumtable''')
        total = cursor.fetchall()
        total=(str(total))
        total=re.sub("\D", "", total)
        total=int(total)
        #add to document
        p = document.add_paragraph('')
        p.add_run('SOFTWARE TABLE SUMMARY').bold = True
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Version'
        hdr_cells[1].text = 'Total'
        hdr_cells[2].text = 'Percentage'        
        for row in records:
            row_cells = table.add_row().cells
            row_cells[0].text = (row[0])
            row_cells[1].text = str(row[1]) 
            percentage = (str((row[1]/total)*100))
            percentage = re.findall('\d+[.]\d',percentage)
            percentage = percentage[0]
            row_cells[2].text = (percentage+'%')
         
        #SOFTWARE TABLE
        #sql query
        cursor.execute('''SELECT devicename, model, iosversion, uptime, confreg FROM swtable''')
        records = cursor.fetchall()
        #add to document
        p = document.add_paragraph('')
        p.add_run('SOFTWARE TABLE').bold = True
        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Model'
        hdr_cells[2].text = 'IOS Version'
        hdr_cells[3].text = 'Uptime'
        hdr_cells[4].text = 'Config Register'
        for row in records:
            row_cells = table.add_row().cells
            row_cells[0].text = (row[0])
            row_cells[1].text = (row[1])
            row_cells[2].text = (row[2])
            row_cells[3].text = (row[3])
            row_cells[4].text = (row[4])           
        
        #HARDWARE SUMMARY TABLE
        #sql query
        cursor.execute('''SELECT model, COUNT(*) FROM hwsumtable GROUP BY model''')
        records = cursor.fetchall()
        cursor.execute('''SELECT COUNT(model) FROM hwsumtable''')
        total = cursor.fetchall()
        total=(str(total))
        total=re.sub("\D", "", total)
        total=int(total)
        #add to document
        p = document.add_paragraph('')
        p.add_run('HARDWARE TABLE SUMMARY').bold = True
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Model'
        hdr_cells[1].text = 'Total'
        hdr_cells[2].text = 'Percentage'        
        for row in records:
            row_cells = table.add_row().cells
            row_cells[0].text = (row[0])
            row_cells[1].text = str(row[1])
            percentage = (str((row[1]/total)*100))
            percentage = re.findall('\d+[.]\d',percentage)
            percentage = percentage[0]
            row_cells[2].text = (percentage+'%')
         
        #HARDWARE CARD TABLE
        #sql query
        cursor.execute('''SELECT devicename, model, card, sn FROM hwcardtable''')
        records = cursor.fetchall()
        p = document.add_paragraph('')
        p.add_run('HARDWARE CARD TABLE').bold = True

        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Model'
        hdr_cells[2].text = 'Card'
        hdr_cells[3].text = 'Slot'
        hdr_cells[4].text = 'Serial Number'

        for row in records:

            row_cells = table.add_row().cells
            row_cells[0].text = (row[0])
            row_cells[1].text = (row[1])
            row_cells[2].text = (row[2])
            row_cells[3].text = ('')
            row_cells[4].text = (row[3])
         
        #CPU SUMMARY TABLE
        #sql query
        cursor.execute('''SELECT devicename, model, process, interrupt, total, status FROM cpusumtable''')
        records = cursor.fetchall()
        p = document.add_paragraph('')
        p.add_run('CPU SUMMARY TABLE').bold = True

        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Model'
        hdr_cells[2].text = 'Total'
        hdr_cells[3].text = 'Process'
        hdr_cells[4].text = 'Interrupt'
        hdr_cells[5].text = 'Status'

        for row in records:

            row_cells = table.add_row().cells
            row_cells[0].text = (row[0])
            row_cells[1].text = (row[1])
            row_cells[2].text = (row[4]+"%")
            row_cells[3].text = (row[2]+"%")
            row_cells[4].text = (row[3]+"%")
            row_cells[5].text = (row[5])
        
        #MEMORY SUMMARY TABLE
        #sql query
        cursor.execute('''SELECT devicename, model, utils, topproc, status FROM memsumtable''')
        records = cursor.fetchall()
        p = document.add_paragraph('')
        p.add_run('MEMORY SUMMARY TABLE').bold = True

        table = document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Model'
        hdr_cells[2].text = 'Processor Memory Utilization'
        hdr_cells[3].text = 'Top Process'
        hdr_cells[4].text = 'Status'

        for row in records:

            row_cells = table.add_row().cells
            row_cells[0].text = (row[0])
            row_cells[1].text = (row[1])
            row_cells[2].text = (row[2]+"%")
            row_cells[3].text = (row[3])
            row_cells[4].text = (row[4])
        
        #SHOW ENVIRONMENT
        #sql query
        cursor.execute('''SELECT * FROM envtable''')
        records = cursor.fetchall()
        #add to document
        p = document.add_paragraph('')
        p.add_run('Hardware Condition Analysis').bold = True
        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'System'
        hdr_cells[2].text = 'Item'   
        hdr_cells[3].text = 'Status'       
        for row in records:
            row_cells = table.add_row().cells
            row_cells[0].text = (row[1])
            row_cells[1].text = (row[2]) 
            row_cells[2].text = (row[3])
            row_cells[3].text = (row[4])

        #LOG
        #sql query
        cursor.execute('''SELECT * FROM logtable''')
        records = cursor.fetchall()
        #add to document
        p = document.add_paragraph('')
        p.add_run('LOG TABLE CHECKING').bold = True
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Device Name'
        hdr_cells[1].text = 'Model'
        hdr_cells[2].text = 'Script'
        for row in records:
            row_cells = table.add_row().cells
            row_cells[0].text = (row[1])
            row_cells[1].text = (row[2])
            row_cells[2].text = (row[3])
        #close database
        db.close()
        
        #save document
        print('Saving Document')
        document.save('preventive_maintenance.docx')
        print('Document has been saved to preventive_maintenance.docx')