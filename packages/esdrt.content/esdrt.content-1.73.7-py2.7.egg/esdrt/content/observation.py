try:
    from cStringIO import StringIO
except ImportError:
    from StringIO import StringIO
import datetime
import re
from itertools import chain
from time import time

from Acquisition import aq_inner
from Acquisition import aq_parent
from z3c.form import button
from z3c.form import field
from z3c.form import interfaces
from z3c.form.browser.checkbox import CheckBoxFieldWidget
from z3c.form.form import Form
from z3c.form.interfaces import ActionExecutionError
from zope import schema
from zope.browsermenu.menu import getMenu
from zope.component import getUtility
from zope.i18n import translate
from zope.interface import Invalid
from zope.interface import alsoProvides
from zope.lifecycleevent.interfaces import IObjectAddedEvent
from zope.lifecycleevent.interfaces import IObjectModifiedEvent
from zope.schema.interfaces import IVocabularyFactory

from Products.statusmessages.interfaces import IStatusMessage

from Products.CMFCore.utils import getToolByName
from Products.CMFEditions import CMFEditionsMessageFactory as _CMFE
from Products.CMFPlone.utils import safe_unicode

from plone import api
from plone.app.contentlisting.interfaces import IContentListing
from plone.app.dexterity.behaviors.discussion import IAllowDiscussion
from plone.app.discussion.interfaces import IConversation
from plone.directives import dexterity
from plone.directives import form
from plone.directives.form import default_value
from plone.memoize import instance
from plone.namedfile.interfaces import IImageScaleTraversable
from plone.z3cform.interfaces import IWrappedForm

from AccessControl import getSecurityManager
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from eea.cache import cache
from esdrt.content import MessageFactory as _
from esdrt.content.constants import ROLE_MSE
from esdrt.content.constants import ROLE_SE
from esdrt.content.constants import ROLE_QE
from esdrt.content.constants import ROLE_LR
from esdrt.content.constants import ROLE_RP1
from esdrt.content.constants import ROLE_RP2
from esdrt.content.constants import ROLE_RE
from esdrt.content.roles.localrolesubscriber import grant_local_roles
from esdrt.content.subscriptions.interfaces import INotificationUnsubscriptions
from esdrt.content.utilities.ms_user import IUserIsMS
from five import grok

from .comment import IComment
from .commentanswer import ICommentAnswer
from .conclusion import IConclusion
from .crf_code_matching import get_category_ldap_from_crf_code
from .crf_code_matching import get_category_value_from_crf_code

HIDDEN_ACTIONS = [
    "/content_status_history",
    "/placeful_workflow_configuration",
]


def hidden(menuitem):
    for action in HIDDEN_ACTIONS:
        if menuitem.get("action").endswith(action):
            return True
    return False


# Cache helper methods
def _user_name(fun, self, userid):
    return (userid, time() // 86400)


# Interface class; used to define content-type schema.
class IObservation(form.Schema, IImageScaleTraversable):
    """
    New review observation
    """

    text = schema.Text(
        title=u"Short description by expert/reviewer",
        required=True,
        description=u"Describe the issue identified. Keep it short, you "
        u"cannot change this description once you have sent it "
        u"to the QE/LR. MS can only see the question once it "
        u"has been approved and sent by the QE/LR. The "
        u"question to the MS should be asked in the Q&A "
        u"tab, not here.",
    )

    country = schema.Choice(
        title=u"Country",
        vocabulary="esdrt.content.eea_member_states",
        required=True,
    )

    crf_code = schema.Choice(
        title=u"CRF category codes",
        vocabulary="esdrt.content.crf_code",
        required=True,
    )

    year = schema.TextLine(
        title=u"Inventory year",
        description=u"Inventory year is the year, a range or a list "
        u"of years or a (e.g. '2012', '2009-2012', "
        u"'2009, 2012, 2013') when the emissions had "
        u"occured for which an issue was observed in the review.",
        required=True,
    )

    form.widget(gas=CheckBoxFieldWidget)
    gas = schema.List(
        title=u"Gas",
        value_type=schema.Choice(vocabulary="esdrt.content.gas",),
        required=True,
    )

    review_year = schema.Int(
        title=u"Review year",
        description=u"Review year is the year in which the inventory was "
        u"submitted and the review was carried out",
        required=True,
    )

    fuel = schema.Choice(
        title=u"Fuel", vocabulary="esdrt.content.fuel", required=False,
    )

    # ghg_source_category = schema.Choice(
    #     title=_(u"CRF category group"),
    #     vocabulary='esdrt.content.ghg_source_category',
    #     required=True,
    # )

    # ghg_source_sectors = schema.Choice(
    #     title=_(u"CRF Sector"),
    #     vocabulary='esdrt.content.ghg_source_sectors',
    #     required=True,
    # )

    ms_key_catagory = schema.Bool(title=u"MS key category",)

    eu_key_catagory = schema.Bool(title=u"EU key category",)

    form.widget(parameter=CheckBoxFieldWidget)
    parameter = schema.List(
        title=u"Parameter",
        value_type=schema.Choice(
            vocabulary="esdrt.content.parameter", required=True,
        ),
        required=True,
    )

    form.widget(highlight=CheckBoxFieldWidget)
    highlight = schema.List(
        title=u"Description flags",
        description=u"Description flags highlight important information "
        u"that is closely related to the main purpose of "
        u"'initial checks' and ESD review",
        value_type=schema.Choice(vocabulary="esdrt.content.highlight",),
        required=False,
        default=[],
    )

    form.write_permission(closing_comments="cmf.ManagePortal")
    closing_comments = schema.Text(
        title=u"Finish request comments", required=False,
    )

    form.write_permission(closing_deny_comments="cmf.ManagePortal")
    closing_deny_comments = schema.Text(
        title=u"Finish deny comments", required=False,
    )

    form.write_permission(closing_comments_phase2="cmf.ManagePortal")
    closing_comments_phase2 = schema.Text(
        title=u"Finish request comments for phase 2", required=False,
    )

    form.write_permission(closing_deny_comments_phase2="cmf.ManagePortal")
    closing_deny_comments_phase2 = schema.Text(
        title=u"Finish deny comments for phase 2", required=False,
    )


@form.validator(field=IObservation["parameter"])
def check_parameter(value):
    if len(value) == 0:
        raise Invalid(u"You need to select at least one parameter")


@form.validator(field=IObservation["gas"])
def check_gas(value):
    if len(value) == 0:
        raise Invalid(u"You need to select at least one gas")


@form.validator(field=IObservation["crf_code"])
def check_crf_code(value):
    """ Check if the user is in one of the group of users
        allowed to add this category CRF Code observations
    """
    category = get_category_ldap_from_crf_code(value)
    user = api.user.get_current()
    groups = user.getGroups()
    valid = False
    for group in groups:
        if group.startswith("extranet-esd-ghginv-sr-%s-" % category):
            valid = True
        if group.startswith("extranet-esd-esdreview-reviewexp-%s-" % category):
            valid = True

    if not valid:
        raise Invalid(
            u"You are not allowed to add observations for this sector category"
        )


@form.validator(field=IObservation["country"])
def check_country(value):
    user = api.user.get_current()
    groups = user.getGroups()
    valid = False
    for group in groups:
        if group.startswith("extranet-esd-ghginv-sr-") and group.endswith(
            "-%s" % value
        ):
            valid = True
        if group.startswith(
            "extranet-esd-esdreview-reviewexp-"
        ) and group.endswith("-%s" % value):
            valid = True

    if not valid:
        raise Invalid(
            u"You are not allowed to add observations for this country"
        )


@form.validator(field=IObservation["year"])
def inventory_year(value):
    """
    Inventory year can be a given year (2014), a range of years (2012-2014)
    or a list of the years (2012, 2014, 2016)
    """

    def string_analyzer(string, separator):
        """ return True is each element of string are Integer
            return False otherwise
        """
        for item in value.split(separator):
            try:
                _ = int(item.strip())
            except ValueError:
                return False
        return True

    try:
        _ = int(value)
        valid = True
    except ValueError:
        # Let's see if it's a range of years or a list of year:
        if "-" in value:
            valid = string_analyzer(value, "-")
        elif "," in value:
            valid = string_analyzer(value, ",")
        else:
            valid = string_analyzer(value, ";")

    if not valid:
        raise Invalid(u"Inventory year format is not correct. ")


@default_value(field=IObservation["review_year"])
def default_year(data):
    return datetime.datetime.now().year


@grok.subscribe(IObservation, IObjectAddedEvent)
@grok.subscribe(IObservation, IObjectModifiedEvent)
def set_title_to_observation(object, event):
    sector = safe_unicode(object.ghg_source_category_value())
    gas = safe_unicode(object.gas_value())
    inventory_year = safe_unicode(str(object.year))
    parameter = safe_unicode(object.parameter_value())
    object.title = u" ".join([sector, gas, inventory_year, parameter])
    grant_local_roles(object)


@grok.subscribe(IObservation, IObjectAddedEvent)
def add_observation(context, event):
    """ When adding an observation, go directly to
        'open' status on the observation
    """
    observation = context
    review_folder = aq_parent(observation)
    with api.env.adopt_roles(roles=["Manager"]):
        if api.content.get_state(obj=review_folder) == "ongoing-review-phase2":
            api.content.transition(obj=observation, transition="go-to-phase2")
            return
        elif api.content.get_state(obj=observation) == "phase1-draft":
            api.content.transition(obj=observation, transition="phase1-approve")


class Observation(dexterity.Container):
    grok.implements(IObservation)
    # Add your class methods and properties here

    def get_values(self):
        """
        Memoized version of values, to speed-up
        """
        return self.values()

    def get_values_cat(self, portal_type=None):
        if portal_type is not None:
            return self.listFolderContents(
                contentFilter={"portal_type": portal_type}
            )
        else:
            return self.listFolderContents()

    def get_crf_code(self):
        """ stupid method to avoid name-clashes with the existing
        vocabularies when cataloging """
        return self.crf_code

    def get_ghg_source_sectors(self):
        """ stupid method to avoid name-clashes with the existing
        vocabularies when cataloging """
        return self.ghg_source_sectors_value()

    def get_highlight(self):
        """ stupid method to avoid name-clashes with the existing
        vocabularies when cataloging """
        return self.highlight

    def country_value(self):
        return self._vocabulary_value(
            "esdrt.content.eea_member_states", self.country
        )

    def crf_code_value(self):
        return self._vocabulary_value("esdrt.content.crf_code", self.crf_code)

    def ghg_source_category_value(self):
        # Get the value of the sector to be used on the LDAP mapping
        return get_category_ldap_from_crf_code(self.crf_code)

    def ghg_source_sectors_value(self):
        # Get the value of the sector to be used
        # on the Observation Metadata screen
        return get_category_value_from_crf_code(self.crf_code)

    def parameter_value(self):
        parameters = [
            self._vocabulary_value("esdrt.content.parameter", p)
            for p in self.parameter
        ]
        return u", ".join(parameters)

    def gas_value(self):
        gases = [
            self._vocabulary_value("esdrt.content.gas", g) for g in self.gas
        ]

        return u", ".join(gases)

    def highlight_value(self):
        if self.highlight:
            highlight = [
                self._vocabulary_value("esdrt.content.highlight", h)
                for h in self.highlight
            ]
            return u", ".join(highlight)
        return u""

    def finish_reason_value(self):
        return self._vocabulary_value(
            "esdrt.content.finishobservationreasons", self.closing_reason
        )

    def finish_deny_reason_value(self):
        return self._vocabulary_value(
            "esdrt.content.finishobservationdenyreasons",
            self.closing_deny_reason,
        )

    def _vocabulary_value(self, vocabulary, term):
        vocab_factory = getUtility(IVocabularyFactory, name=vocabulary)
        vocabulary = vocab_factory(self)
        try:
            value = vocabulary.getTerm(term)
            return value.title
        except LookupError:
            return u""

    def get_status(self):
        return api.content.get_state(self)

    def can_draft_conclusions(self):
        questions = self.get_values_cat("Question")
        if len(questions) > 0:
            q = questions[0]
            return q.get_state_api() in [
                "phase1-draft",
                "phase1-drafted",
                "phase1-recalled-lr",
                "phase1-closed",
                "phase2-draft",
                "phase2-drafted",
                "phase2-recalled-lr",
                "phase2-closed",
            ]
        else:
            return True

    def can_close(self):
        if self.get_status() in ["phase1-pending", "phase2-pending"]:
            questions = self.get_values_cat("Question")
            if len(questions) > 0:
                for q in questions:
                    if q.get_state_api() not in [
                        "phase1-closed",
                        "phase2-closed",
                    ]:
                        return False
                return True

        return False

    def wf_location(self):
        if self.get_status() == "phase1-draft":
            return "Sector expert"
        elif self.get_status() == "phase2-draft":
            return "Review expert"
        elif self.get_status() == "phase1-closed":
            return "Quality expert"
        elif self.get_status() == "phase2-closed":
            return "Lead reviewer"
        elif self.get_status() == "phase1-conclusions":
            return "Sector expert"
        elif self.get_status() == "phase2-conclusions":
            return "Review expert"
        elif self.get_status() in [
            "phase1-conclusion-discussion",
            "phase2-conclusion-discussion",
        ]:
            return "Counterpart"
        elif self.get_status() == "phase1-close-requested":
            return "Quality expert"
        elif self.get_status() == "phase2-close-requested":
            return "Lead reviewer"
        else:
            questions = self.get_values_cat("Question")
            if questions:
                question = questions[0]
                state = question.get_state_api()
                if state in ["phase1-draft", "phase1-closed"]:
                    return "Sector expert"
                if state in ["phase2-draft", "phase2-closed"]:
                    return "Review expert"
                elif state in [
                    "phase1-counterpart-comments",
                    "phase2-counterpart-comments",
                ]:
                    return "Counterparts"
                elif state in ["phase1-drafted", "phase1-recalled-lr"]:
                    return "Quality Expert"
                elif state in ["phase2-drafted", "phase2-recalled-lr"]:
                    return "Lead reviewer"
                elif state in [
                    "phase1-pending",
                    "phase1-answered",
                    "phase1-pending-answer-drafting",
                    "phase1-recalled-msa",
                    "phase2-pending",
                    "phase2-answered",
                    "phase2-pending-answer-drafting",
                    "phase2-recalled-msa",
                ]:
                    return "Member state coordinator"
                elif state in [
                    "phase1-expert-comments",
                    "phase2-expert-comments",
                ]:
                    return "Member state experts"
            else:
                if self.get_status().startswith("phase1"):
                    return "Sector expert"
                else:
                    return "Review expert"

    def wf_status(self):
        if self.get_status() in ["phase1-draft", "phase2-draft"]:
            return ["Observation created", "observationIcon"]
        elif self.get_status() in ["phase1-closed", "phase2-closed"]:
            return ["Observation finished", "observationIcon"]
        elif self.get_status() in [
            "phase1-close-requested",
            "phase2-close-requested",
        ]:
            return ["Observation finish requested", "observationIcon"]
        elif self.get_status() in ["phase1-conclusions", "phase2-conclusions"]:
            return ["Conclusion ongoing", "conclusionIcon"]
        elif self.get_status() in [
            "phase1-conclusion-discussion",
            "phase2-conclusion-discussion",
        ]:
            return ["Counterparts comments requested", "conclusionIcon"]
        else:
            questions = self.get_values_cat("Question")
            if questions:
                question = questions[-1]
                state = question.get_state_api()
                if state in ["phase1-draft", "phase2-draft"]:
                    return ["Question drafted", "questionIcon"]
                elif state in [
                    "phase1-counterpart-comments",
                    "phase2-counterpart-comments",
                ]:
                    return ["Counterpart's comments requested", "questionIcon"]
                elif state in ["phase1-answered", "phase2-answered"]:
                    return ["Pending question", "questionIcon"]
                elif state in [
                    "phase1-pending",
                    "phase1-pending-answer-drafting",
                    "phase1-recalled-msa",
                    "phase2-pending",
                    "phase2-pending-answer-drafting",
                    "phase2-recalled-msa",
                ]:
                    return ["Open question", "questionIcon"]
                elif state in [
                    "phase1-drafted",
                    "phase1-recalled-lr",
                    "phase2-drafted",
                    "phase2-recalled-lr",
                ]:
                    return ["Draft question", "questionIcon"]
                elif state in [
                    "phase1-expert-comments",
                    "phase2-expert-comments",
                ]:
                    return ["MS expert comments requested", "questionIcon"]
                elif state in ["phase1-closed", "phase2-closed"]:
                    return ["Closed question", "questionIcon"]
            else:
                return ["Observation created", "observationIcon"]

        return ["Unknown", "observationIcon"]

    def observation_status(self):
        status = self.observation_question_status()
        my_status = self.get_status()

        if status in [
            "phase1-draft",
            "phase2-draft",
            "phase1-counterpart-comments",
            "phase2-counterpart-comments",
            "observation-phase1-draft",
            "observation-phase2-draft",
        ]:
            return "SRRE"
        elif status in [
            "phase1-drafted",
            "phase2-drafted",
            "phase1-recalled-lr",
            "phase2-recalled-lr",
        ]:
            return "LRQE"
        elif status in [
            "phase1-pending",
            "phase2-pending",
            "phase1-pending-answer-drafting",
            "phase2-pending-answer-drafting",
            "phase1-expert-comments",
            "phase2-expert-comments",
            "phase1-recalled-msa",
            "phase1-recalled-msa",
        ]:
            return "MSC"
        elif status in ["phase1-answered", "phase2-answered"]:
            return "answered"
        elif status in [
            "phase1-conclusions",
            "phase2-conclusions",
            "phase1-conclusion-discussion",
            "phase2-conclusion-discussion",
        ]:
            return "conclusions"
        elif status in ["phase1-close-requested", "phase2-close-requested"]:
            return "close-requested"
        elif status in ["phase1-closed", "phase2-closed"]:
            if status == "phase1-closed":
                conclusion = self.get_conclusion()
                conclusion_reason = (
                    conclusion and conclusion.closing_reason or " "
                )
                if conclusion_reason == "no-conclusion-yet":
                    return "SRRE"
                elif not my_status.endswith("-closed"):
                    return "answered"
                else:
                    return "finalised"
            elif status == "phase2-closed":
                conclusion = self.get_conclusion_phase2()
                conclusion_reason = (
                    conclusion and conclusion.closing_reason or " "
                )
                if conclusion_reason == "no-conclusion-yet":
                    return "SRRE"
                elif not my_status.endswith("-closed"):
                    return "answered"
                else:
                    return "finalised"
        else:
            return status

    def overview_status(self):
        status = self.get_status()
        closed_val = "closed ({reason})"
        if status == "phase1-closed":
            conclusion = self.get_conclusion()
            if conclusion:
                return closed_val.format(reason=conclusion.reason_value())

        elif status == "phase2-closed":
            conclusion = self.get_conclusion_phase2()
            if conclusion:
                return closed_val.format(reason=conclusion.reason_value())

        else:
            return "open"

    @staticmethod
    def is_secretariat():
        user = api.user.get_current()
        return "Manager" in user.getRoles()

    @cache(_user_name)
    def _author_name(self, userid):
        if userid:
            user = api.user.get(username=userid)
            if user:
                return user.getProperty("fullname", userid)

        return userid

    def get_author_name(self, userid=None):
        if not userid:
            userid = self.Creator()

        return self._author_name(userid)

    def myHistory(self):
        observation_history = self.workflow_history.get(
            "esd-review-workflow", []
        )
        observation_wf = []
        question_wf = []
        for item in observation_history:
            item["role"] = item["actor"]
            item["object"] = "observationIcon"
            item["author"] = self.get_author_name(item["actor"])
            if item["review_state"] == "phase1-draft":
                item["state"] = "Draft observation"
                if self.observation_phase() == "phase1-observation":
                    item["role"] = "Sector expert"
                else:
                    item["role"] = "Review expert"
                observation_wf.append(item)
            elif (
                item["review_state"] == "phase1-pending"
                and item["action"] == "phase1-approve"
            ):
                item["state"] = "Pending"
                # Do not add
            elif (
                item["review_state"] == "phase1-pending"
                and item["action"] == "phase1-reopen"
            ):
                item["state"] = "Observation reopened"
                item["role"] = "Sector expert"
                observation_wf.append(item)
            elif item["review_state"] == "phase1-closed":
                item["state"] = "Closed observation"
                item["role"] = "Quality expert"
                observation_wf.append(item)
            elif item["review_state"] == "phase1-close-requested":
                item["state"] = "Finalisation requested"
                item["role"] = "Sector expert"
                observation_wf.append(item)
            elif (
                item["review_state"] == "phase1-conclusions"
                and item["action"] == "phase1-deny-closure"
            ):
                item["state"] = "Finalisation denied"
                item["role"] = "Quality expert"
                observation_wf.append(item)
            elif item["review_state"] == "phase1-conclusion-discussion":
                item["state"] = "Conclusion comments requested"
                item["role"] = "Sector expert"
                item["object"] = "conclusionIcon"
                observation_wf.append(item)
            elif (
                item["review_state"] == "phase1-conclusions"
                and item["action"] == "phase1-finish-comments"
            ):
                item["state"] = "Conclusion comments closed"
                item["role"] = "Sector expert"
                item["object"] = "conclusionIcon"
                observation_wf.append(item)
            elif (
                item["review_state"] == "phase1-conclusions"
                and item["action"] == "phase1-draft-conclusions"
            ):
                item["state"] = "Conclusion drafting"
                item["role"] = "Sector expert"
                item["object"] = "conclusionIcon"
                observation_wf.append(item)
            elif item["action"] == "phase1-send-to-team-2":
                item["state"] = "Handed to phase 2"
                item["role"] = "Quality expert"
                observation_wf.append(item)
            elif item["review_state"] == "phase2-draft":
                item["state"] = "Draft observation"
                item["role"] = "Review expert"
                observation_wf.append(item)
            elif (
                item["review_state"] == "phase2-pending"
                and item["action"] == "phase2-approve"
            ):
                item["state"] = "Pending"
                # Do not add
            elif (
                item["review_state"] == "phase2-pending"
                and item["action"] == "phase2-reopen"
            ):
                item["state"] = "Observation reopened"
                item["role"] = "Review expert"
                observation_wf.append(item)
            elif item["review_state"] == "phase2-pending":
                item["state"] = "Pending"
                # Do not add
            elif item["review_state"] == "phase2-closed":
                item["state"] = "Closed observation"
                item["role"] = "Lead Reviewer"
                observation_wf.append(item)
            elif item["review_state"] == "phase2-close-requested":
                item["state"] = "Finalisation requested"
                item["role"] = "Review expert"
                observation_wf.append(item)
            elif (
                item["review_state"] == "phase2-conclusions"
                and item["action"] == "phase2-deny-finishing-observation"
            ):
                item["state"] = "Finalisation denied"
                item["role"] = "Lead reviewer"
                observation_wf.append(item)
            elif item["review_state"] == "phase2-conclusion-discussion":
                item["state"] = "Conclusion comments requested"
                item["role"] = "Review expert"
                item["object"] = "conclusionIcon"
                observation_wf.append(item)
            elif (
                item["review_state"] == "phase2-conclusions"
                and item["action"] == "phase2-finish-comments"
            ):
                item["state"] = "Conclusion comments closed"
                item["role"] = "Review expert"
                item["object"] = "conclusionIcon"
                observation_wf.append(item)
            elif (
                item["review_state"] == "phase2-conclusions"
                and item["action"] == "phase2-draft-conclusions"
            ):
                item["state"] = "Conclusion drafting"
                item["role"] = "Review expert"
                item["object"] = "conclusionIcon"
                observation_wf.append(item)
            else:
                item["state"] = "*" + item["review_state"] + "*"
                observation_wf.append(item)

        history = list(observation_wf)
        questions = self.get_values_cat()

        if questions:
            question = questions[0]
            question_history = question.workflow_history.get(
                "esd-question-review-workflow", []
            )
            for item in question_history:
                item["role"] = item["actor"]
                item["object"] = "questionIcon"
                item["author"] = self.get_author_name(item["actor"])
                if (
                    item["review_state"] == "phase1-draft"
                    and item["action"] is None
                ):
                    item["state"] = "Draft question"
                    if self.observation_phase() == "phase1-observation":
                        item["role"] = "Sector expert"
                    else:
                        item["role"] = "Quality expert"
                    question_wf.append(item)
                elif item["review_state"] == "phase1-counterpart-comments":
                    item["state"] = "Requested counterparts comments"
                    item["role"] = "Sector expert"
                    question_wf.append(item)
                elif (
                    item["review_state"] == "phase1-draft"
                    and item["action"] == "phase1-send-comments"
                ):
                    item["state"] = "Counterparts comments closed"
                    item["role"] = "Sector expert"
                    question_wf.append(item)
                elif item["review_state"] == "phase1-drafted":
                    item["state"] = "Sent to Quality expert"
                    item["role"] = "Sector expert"
                    question_wf.append(item)
                elif (
                    item["review_state"] == "phase1-draft"
                    and item["action"] == "phase1-recall-sre"
                ):
                    item["state"] = "Question recalled"
                    item["role"] = "Sector expert"
                    question_wf.append(item)
                elif (
                    item["review_state"] == "phase1-draft"
                    and item["action"] == "phase1-redraft"
                ):
                    item["state"] = "Question redrafted"
                    item["role"] = "Quality expert"
                    question_wf.append(item)
                elif (
                    item["review_state"] == "phase1-pending"
                    and item["action"] == "phase1-approve-question"
                ):
                    item[
                        "state"
                    ] = "Question approved and sent to MS coodinator"
                    item["role"] = "Quality expert"
                    question_wf.append(item)
                elif item["review_state"] == "phase1-recalled-lr":
                    item["state"] = "Question recalled"
                    item["role"] = "Quality expert"
                elif item["review_state"] == "phase1-answered":
                    item["state"] = "Answer sent"
                    item["role"] = "Member state coordinator"
                    question_wf.append(item)
                elif item["review_state"] == "phase1-expert-comments":
                    item["state"] = "MS expert comments requested"
                    item["role"] = "Member state coordinator"
                    question_wf.append(item)
                elif item["review_state"] == "phase1-pending-answer-drafting":
                    item["state"] = "Member state expert comments closed"
                    item["role"] = "Member state coordinator"
                    question_wf.append(item)
                elif item["review_state"] == "phase1-recalled-msa":
                    item["state"] = "Answer recalled"
                    item["role"] = "Member state coordinator"
                    question_wf.append(item)
                elif (
                    item["action"] == "phase1-closed"
                    and item["action"] == "phase1-validate-answer-msa"
                ):
                    item["state"] = "Sector expert"
                    item["role"] = "Answer acknowledged"
                    question_wf.append(item)
                elif (
                    item["review_state"] == "phase1-draft"
                    and item["action"] == "phase1-reopen"
                ):
                    item["state"] = "Reopened"
                    # Do not add
                elif item["review_state"] == "phase1-closed":
                    item["state"] = "Closed"
                    # Do not add
                elif (
                    item["review_state"] == "phase2-draft"
                    and item["action"] == "phase2-reopen"
                ):
                    item["state"] = "Draft question"
                    item["role"] = "Review expert"
                    question_wf.append(item)
                elif item["review_state"] == "phase2-counterpart-comments":
                    item["state"] = "Requested counterparts comments"
                    item["role"] = "Review expert"
                    question_wf.append(item)
                elif (
                    item["review_state"] == "phase2-draft"
                    and item["action"] == "phase2-send-comments"
                ):
                    item["state"] = "Counterparts comments closed"
                    item["role"] = "Review expert"
                    question_wf.append(item)
                elif item["review_state"] == "phase2-drafted":
                    item["state"] = "Sent to LR"
                    item["role"] = "Review expert"
                    question_wf.append(item)
                elif (
                    item["review_state"] == "phase2-draft"
                    and item["action"] == "phase2-recall-sre"
                ):
                    item["state"] = "Question recalled"
                    item["role"] = "Review expert"
                    question_wf.append(item)
                elif (
                    item["review_state"] == "phase2-draft"
                    and item["action"] == "phase2-redraft"
                ):
                    item["state"] = "Question redrafted"
                    item["role"] = "Review expert"
                    question_wf.append(item)
                elif item["review_state"] == "phase2-draft":
                    # Do not add
                    pass
                elif (
                    item["review_state"] == "phase2-pending"
                    and item["action"] == "phase2-approve-question"
                ):
                    item[
                        "state"
                    ] = "Question approved and sent to MS coordinator"
                    item["role"] = "Lead reviewer"
                    question_wf.append(item)
                elif item["review_state"] == "phase2-recalled-lr":
                    item["state"] = "Question recalled"
                    item["role"] = "Lead reviewer"
                elif item["review_state"] == "phase2-answered":
                    item["state"] = "Answer sent"
                    item["role"] = "Member state coordinator"
                    question_wf.append(item)
                elif item["review_state"] == "phase2-expert-comments":
                    item["state"] = "MS expert comments requested"
                    item["role"] = "Member state coordinator"
                    question_wf.append(item)
                elif item["review_state"] == "phase2-pending-answer-drafting":
                    item["state"] = "Member state expert comments closed"
                    item["role"] = "Member state coordinator"
                    question_wf.append(item)
                elif item["review_state"] == "phase2-recalled-msa":
                    item["state"] = "Answer recalled"
                    item["role"] = "Member state coordinator"
                    question_wf.append(item)
                elif (
                    item["action"] == "phase2-validate-answer-msa"
                    and item["action"] == "phase2-validate-answer-msa"
                ):
                    item["state"] = "Review expert"
                    item["role"] = "Answer acknowledged"
                    question_wf.append(item)
                elif (
                    item["review_state"] == "phase2-draft"
                    and item["action"] == "phase2-reopen"
                ):
                    item["state"] = "Reopened"
                    # Do not add
                elif item["review_state"] == "phase2-closed":
                    item["state"] = "Closed"
                    # Do not add
                else:
                    item["state"] = "*" + item["review_state"] + "*"
                    item["role"] = item["actor"]
                    question_wf.append(item)

            history = list(observation_wf) + list(question_wf)

        history.sort(key=lambda x: x["time"], reverse=False)
        return history

    def can_edit(self):
        sm = getSecurityManager()
        return sm.checkPermission("Modify portal content", self)

    def get_question(self):
        questions = self.get_values_cat("Question")

        if questions:
            question = questions[-1]
            return question

    def observation_question_status(self):
        questions = self.get_values_cat("Question")
        if (
            self.get_status() != "phase1-pending"
            and self.get_status() != "phase2-pending"
        ):
            if self.get_status() in [
                "phase2-conclusions",
            ]:
                if questions:
                    question = questions[-1]
                    question_state = api.content.get_state(question)
                    if question_state != "phase2-closed":
                        return question_state
            return self.get_status()
        else:
            if questions:
                question = questions[-1]
                state = api.content.get_state(question)
                return state
            else:
                if self.get_status().startswith("phase1"):
                    return "observation-phase1-draft"
                else:
                    return "observation-phase2-draft"

    def observation_questions_workflow(self):
        questions = self.get_values_cat("Question")
        if not questions:
            return tuple()

        # there is always only one question.
        question = questions[0]

        items = question.values()

        comments = [i for i in items if i.portal_type == "Comment"]
        answers = [i for i in items if i.portal_type == "CommentAnswer"]

        len_comments = len(comments)
        obs_status = self.observation_status()

        return tuple(["Answered"] * (len_comments - 1) + [obs_status])

    def observation_css_class(self):
        if self.highlight:
            if self.get_status().startswith("phase1"):
                if "psi" in self.highlight:
                    return "psiBackground"
            else:
                if self.get_status() == "phase2-closed":
                    con_phase2 = self.get_conclusion_phase2()
                    if con_phase2:
                        if con_phase2.closing_reason == "technical-correction":
                            return "technicalCorrectionBackground"

                elif "ptc" in self.highlight:
                    return "ptcBackground"

    def observation_is_potential_significant_issue(self):
        if self.highlight:
            return "psi" in self.highlight
        return False

    def observation_is_potential_technical_correction(self):
        if self.highlight:
            return "ptc" in self.highlight
        return False

    def observation_is_technical_correction(self):
        if self.get_status() == "phase2-closed":
            con_phase2 = self.get_conclusion_phase2()
            if con_phase2:
                return con_phase2.closing_reason == "technical-correction"
        return False

    def get_conclusion(self):
        conclusions = self.get_values_cat("Conclusion")
        mtool = api.portal.get_tool("portal_membership")
        if conclusions and mtool.checkPermission("View", conclusions[0]):
            return conclusions[0]
        return None

    def get_conclusion_phase2(self):
        conclusions = self.get_values_cat("ConclusionsPhase2")
        mtool = api.portal.get_tool("portal_membership")
        if conclusions and mtool.checkPermission("View", conclusions[0]):
            return conclusions[0]
        return None

    def last_question_reply_number(self):
        questions = self.get_values_cat("Question")
        replynum = 0
        if questions:
            comments = [
                c for c in questions[-1].values() if c.portal_type == "Comment"
            ]
            if comments:
                last = comments[-1]
                disc = IConversation(last)
                return disc.total_comments

        return replynum

    def last_answer_reply_number(self):
        questions = self.get_values_cat("Question")
        replynum = 0
        if questions:
            comments = [
                c
                for c in questions[-1].values()
                if c.portal_type == "CommentAnswer"
            ]
            if comments:
                last = comments[-1]
                disc = IConversation(last)
                return disc.total_comments

        return replynum

    def reply_comments_by_mse(self):
        question = self.get_question()
        commentators = []
        if question:
            commentators = list(
                set(
                    chain.from_iterable(
                        [
                            IConversation(c).commentators
                            for c in question.values()
                        ]
                    )
                )
            )

        return [
            uid
            for uid in commentators
            if ROLE_MSE in api.user.get_roles(username=uid, obj=self)
        ]

    def observation_already_replied(self):

        questions = self.get_values_cat("Question")
        if questions:
            question = questions[0]
            winfo = question.workflow_history
            state = self.get_status()
            for witem in winfo.get("esd-question-review-workflow", []):
                if state.startswith("phase1-"):
                    if witem.get("review_state", "") == "phase1-answered":
                        return True
                elif state.startswith("phase2-"):
                    if witem.get("review_state", "") == "phase2-answered":
                        return True

        return False

    def observation_phase(self):
        status = api.content.get_state(self)
        if status.startswith("phase1-"):
            return "phase1-observation"
        else:
            return "phase2-observation"

    def can_add_followup(self):
        status = self.get_status()
        return status in ["phase1-conclusions", "phase2-conclusions"]


# View class
# The view will automatically use a similarly named template in
# templates called observationview.pt .
# Template filenames should be all lower case.
# The view will render when you request a content object with this
# interface with "/@@view" appended unless specified otherwise
# using grok.name below.
# This will make this view the default view for your content-type

grok.templatedir("templates")


class AddForm(dexterity.AddForm):
    grok.name("esdrt.content.observation")
    grok.context(IObservation)
    grok.require("esdrt.content.AddObservation")

    label = "Observation"
    description = " "

    def updateWidgets(self):
        super(AddForm, self).updateWidgets()
        self.fields["IDublinCore.title"].field.required = False
        self.widgets["IDublinCore.title"].mode = interfaces.HIDDEN_MODE
        self.widgets["IDublinCore.description"].mode = interfaces.HIDDEN_MODE
        self.widgets["text"].rows = 15
        if not Observation.is_secretariat():
            self.widgets["review_year"].readonly = "readonly"
        self.groups = [
            g for g in self.groups if g.label == "label_schema_default"
        ]

    def updateActions(self):
        super(AddForm, self).updateActions()
        self.actions["save"].title = u"Save Observation"
        self.actions["save"].addClass("defaultWFButton")
        self.actions["cancel"].title = u"Delete Observation"

        for k in self.actions.keys():
            self.actions[k].addClass("standardButton")


class ObservationMixin(grok.View):
    grok.baseclass()

    @property
    def user_roles(self):
        user = api.user.get_current()
        return api.user.get_roles(username=user.getId(), obj=self.context)

    def is_se(self):
        return ROLE_SE in self.user_roles

    def is_qe(self):
        return ROLE_QE in self.user_roles

    def can_view_redraft_reason(self):
        allowed_roles = (
            ROLE_LR, ROLE_RP1, ROLE_RP2, ROLE_QE, ROLE_SE, ROLE_RE
        )
        return (
            self.context.is_secretariat()
            or set(self.user_roles).intersection(allowed_roles)
        )

    def wf_info(self):
        context = aq_inner(self.context)
        wf = getToolByName(context, "portal_workflow")
        comments = wf.getInfoFor(
            self.context, "comments", wf_id="esd-review-workflow"
        )
        actor = wf.getInfoFor(
            self.context, "actor", wf_id="esd-review-workflow"
        )
        tim = wf.getInfoFor(self.context, "time", wf_id="esd-review-workflow")
        return {"comments": comments, "actor": actor, "time": tim}

    def isManager(self):
        sm = getSecurityManager()
        context = aq_inner(self.context)
        return sm.checkPermission("Manage portal", context)

    def get_user_name(self, userid):
        # Check users roles
        country = self.context.country_value()
        sector = self.context.ghg_source_sectors_value()
        return " - ".join([country, sector])

    def get_menu_actions(self):
        context = aq_inner(self.context)
        menu_items = getMenu(
            "plone_contentmenu_workflow", context, self.request
        )
        return [mitem for mitem in menu_items if not hidden(mitem)]

    def get_questions(self):
        return IContentListing(self.context.get_values_cat("Question"))

    def can_delete_observation(self):
        is_draft = self.context.get_status() in [
            "phase1-pending",
            "phase2-pending",
        ]
        questions = len(self.context.get_values_cat("Question"))
        # If observation has conclusion cannot be deleted (Ticket #26992)
        conclusions = len(self.context.get_values_cat("Conclusion"))
        return is_draft and not questions and not conclusions

    def can_add_question(self):
        sm = getSecurityManager()
        questions = len(self.context.get_values_cat("Question"))
        return (
            sm.checkPermission("esdrt.content: Add Question", self)
            and not questions
        )

    def can_edit(self):
        sm = getSecurityManager()
        # If observation has conclusion cannot be edited (Ticket #26992)
        conclusions = len(self.context.get_values_cat("Conclusion"))
        return (
            sm.checkPermission("Modify portal content", self.context)
            and not conclusions
        )

    def get_conclusion(self):
        sm = getSecurityManager()
        conclusions = self.context.get_values_cat("Conclusion")
        if conclusions and sm.checkPermission("View", conclusions[0]):
            return conclusions[0]

        return None

    def get_conclusion_phase2(self):
        sm = getSecurityManager()
        conclusions = self.context.get_values_cat("ConclusionsPhase2")
        if conclusions and sm.checkPermission("View", conclusions[0]):
            return conclusions[0]

        return None

    def existing_conclusion(self):
        status = self.context.get_status()
        if status.startswith("phase1-"):
            conclusion = self.get_conclusion()
        else:
            conclusion = self.get_conclusion_phase2()

        return conclusion and True or False

    def can_add_conclusion(self):
        sm = getSecurityManager()
        status = self.context.get_status()
        if status.startswith("phase1-"):
            conclusion = self.get_conclusion()
            return (
                sm.checkPermission(
                    "esdrt.content: Add Conclusion", self.context
                )
                and not conclusion
            )
        else:
            conclusion = self.get_conclusion_phase2()
            return (
                sm.checkPermission(
                    "esdrt.content: Add ConclusionsPhase2", self.context
                )
                and not conclusion
            )

    def show_description(self):
        questions = self.get_questions()
        sm = getSecurityManager()
        if questions:
            question = questions[-1]
            return sm.checkPermission("View", question.getObject())
        else:
            return not getUtility(IUserIsMS)(self.context)

    def show_internal_notes(self):
        return not getUtility(IUserIsMS)(self.context)

    def _can_view_conclusion_remarks(self, state):
        result = True
        user_is_ms = getUtility(IUserIsMS)(self.context)
        if self.context.is_secretariat():
            result = True
        elif user_is_ms:
            result = self.context.get_status() == state
        return result

    def can_view_conclusion_remarks_phase1(self):
        return self._can_view_conclusion_remarks("phase1-closed")

    def can_view_conclusion_remarks_phase2(self):
        return self._can_view_conclusion_remarks("phase2-closed")

    def add_question_form(self):
        from plone.z3cform.interfaces import IWrappedForm

        form_instance = AddQuestionForm(self.context, self.request)
        alsoProvides(form_instance, IWrappedForm)
        return form_instance()

    def has_local_notifications_settings(self):
        user = api.user.get_current()
        adapted = INotificationUnsubscriptions(self.context)
        data = adapted.get_user_data(user.getId())

        return data and True or False

    # Question view
    def question(self):
        questions = self.get_questions()
        if questions:
            return questions[0].getObject()

    def get_chat(self):
        sm = getSecurityManager()
        question = self.question()
        if question:
            values = [
                v for v in question.values() if sm.checkPermission("View", v)
            ]
            # return question.values()
            return values

    def is_old_qa(self, comment):
        return comment.creation_date.year() < datetime.datetime.now().year

    def actions(self):
        context = aq_inner(self.context)
        question = self.question()
        observation_menu_items = getMenu(
            "plone_contentmenu_workflow", context, self.request
        )
        menu_items = observation_menu_items
        if question:
            question_menu_items = getMenu(
                "plone_contentmenu_workflow", question, self.request
            )

            menu_items = question_menu_items + observation_menu_items
        return [mitem for mitem in menu_items if not hidden(mitem)]

    def get_user_name(self, userid, question=None):
        # check users
        if question is not None:
            country = self.context.country_value()
            sector = self.context.ghg_source_sectors_value()
            if question.portal_type == "Comment":
                return " - ".join([country, sector])
            elif question.portal_type == "CommentAnswer":
                return " - ".join([country, "Coordinator"])

        if userid:
            user = api.user.get(username=userid)
            return user.getProperty("fullname", userid)
        return ""

    def can_add_comment(self):
        sm = getSecurityManager()
        question = self.question()
        if question:
            permission = sm.checkPermission(
                "esdrt.content: Add Comment", question
            )
            questions = [
                q for q in question.values() if q.portal_type == "Comment"
            ]
            answers = [
                q for q in question.values() if q.portal_type == "CommentAnswer"
            ]
            obs_state = self.context.get_status()
            return (
                permission
                and len(questions) == len(answers)
                and obs_state != "phase1-closed"
            )
        else:
            return False

    def can_add_answer(self):
        sm = getSecurityManager()
        question = self.question()
        if question:
            permission = sm.checkPermission(
                "esdrt.content: Add CommentAnswer", question
            )
            questions = [
                q for q in question.values() if q.portal_type == "Comment"
            ]
            answers = [
                q for q in question.values() if q.portal_type == "CommentAnswer"
            ]
            return permission and len(questions) > len(answers)
        else:
            return False

    def add_answer_form(self):
        form_instance = AddAnswerForm(self.context, self.request)
        alsoProvides(form_instance, IWrappedForm)
        return form_instance()

    def can_see_comments(self):
        state = self.question().get_state_api()
        # return state in ['draft', 'counterpart-comments', 'drafted']
        return state in ["phase1-counterpart-comments"]

    def add_comment_form(self):
        form_instance = AddCommentForm(self.context, self.request)
        alsoProvides(form_instance, IWrappedForm)
        return form_instance()

    def add_conclusion_form(self):
        form_instance = AddConclusionForm(self.context, self.request)
        alsoProvides(form_instance, IWrappedForm)
        return form_instance()

    def in_conclusions_phase1(self):
        state = self.context.get_status()
        return state in [
            "phase1-conclusions",
            "phase1-conclusion-discussion",
            "phase1-close-requested",
            "phase1-closed",
        ]

    def in_conclusions_phase2(self):
        state = self.context.get_status()
        return state in [
            "phase2-conclusions",
            "phase2-conclusion-discussion",
            "phase2-close-requested",
            "phase2-closed",
        ]

    def get_last_editable_thing(self):
        CONCLUSIONS_PHASE_1 = [
            "phase1-conclusions",
            "phase1-conclusion-discussion",
            "phase1-close-requested",
        ]
        CONCLUSIONS_PHASE_2 = [
            "phase2-conclusions",
            "phase2-conclusion-discussion",
            "phase2-close-requested",
        ]
        MS_OBSERVATION = [
            "phase1-pending",
            "phase2-pending",
        ]

        MS_QUESTION = [
            "phase1-pending",
            "phase1-pending-answer-drafting",
            "phase1-expert-comments",
            "phase2-pending",
            "phase2-pending-answer-drafting",
            "phase2-expert-comments",
        ]
        state = self.context.get_status()
        if state in CONCLUSIONS_PHASE_1:
            return self.context.get_conclusion()
        elif state in CONCLUSIONS_PHASE_2:
            return self.context.get_conclusion_phase2()
        else:
            question = self.question()
            if question is not None:
                qs = question.get_questions()
                if qs:
                    return qs[-1].getObject()

        return None

    def update(self):
        context = self.get_last_editable_thing()
        if context is not None:
            if context.can_edit():
                try:
                    history_metadata = self.repo_tool.getHistoryMetadata(
                        context
                    )
                except:
                    history_metadata = None
                if history_metadata:
                    retrieve = history_metadata.retrieve
                    getId = history_metadata.getVersionId
                    history = self.history = []
                    # Count backwards from most recent to least recent
                    for i in xrange(
                        history_metadata.getLength(countPurged=False) - 1,
                        -1,
                        -1,
                    ):
                        version = retrieve(i, countPurged=False)[
                            "metadata"
                        ].copy()
                        version["version_id"] = getId(i, countPurged=False)
                        history.append(version)
                    dt = getToolByName(self.context, "portal_diff")

                    version1 = self.request.get("one", None)
                    version2 = self.request.get("two", None)

                    if version1 is None and version2 is None:
                        self.history.sort(
                            lambda x, y: cmp(
                                x.get("version_id", ""), y.get("version_id")
                            ),
                            reverse=True,
                        )
                        version1 = self.history[-1].get("version_id", "current")
                        if len(self.history) > 1:
                            version2 = self.history[-2].get(
                                "version_id", "current"
                            )
                        else:
                            version2 = "current"
                    elif version1 is None:
                        version1 = "current"
                    elif version2 is None:
                        version2 = "current"

                    self.request.set("one", version1)
                    self.request.set("two", version2)
                    changeset = dt.createChangeSet(
                        self.getVersion(version2),
                        self.getVersion(version1),
                        id1=self.versionTitle(version2),
                        id2=self.versionTitle(version1),
                    )
                    self.changes = [
                        change
                        for change in changeset.getDiffs()
                        if not change.same
                    ]

    @property
    def repo_tool(self):
        return getToolByName(self.context, "portal_repository")

    def getVersion(self, version):
        context = self.get_last_editable_thing()
        if version == "current":
            return context
        else:
            return self.repo_tool.retrieve(context, int(version)).object

    def versionName(self, version):
        """
        Copied from @@history_view
        Translate the version name. This is needed to allow translation
        when `version` is the string 'current'.
        """
        return _CMFE(version)

    def versionTitle(self, version):
        version_name = self.versionName(version)

        return translate(
            _CMFE(u"version ${version}", mapping=dict(version=version_name)),
            context=self.request,
        )

    def isChatCurrent(self):
        status = api.content.get_state(self.context)
        if status in [
            "phase1-draft",
            "phase1-pending",
            "phase2-draft",
            "phase2-pending",
        ]:
            return True
        else:
            return False


class ObservationView(ObservationMixin):
    grok.context(IObservation)
    grok.require("zope2.View")
    grok.name("view")

    def get_current_counterparters(self):
        """ Return list of current counterparters,
            if the user can see counterpart action
        """
        actions = [action["action"] for action in self.actions()]
        if not any("counterpart_form" in action for action in actions):
            return []

        target = self.context
        local_roles = target.get_local_roles()
        users = [u[0] for u in local_roles if "CounterPart" in u[1]]
        return [api.user.get(user) for user in users]

    def can_export_observation(self):
        sm = getSecurityManager()
        return sm.checkPermission(
            "esdrt.content: Export an Observation", self.context
        )


class DiffedView(ObservationView):
    grok.name("diffedview")
    grok.context(IObservation)
    grok.require("zope2.View")


class ExportAsDocView(ObservationMixin):
    grok.name("export_as_docx")
    grok.context(IObservation)
    grok.require("esdrt.content.ExportAnObservation")

    def strip_special_chars(self, s):
        """ return s without special chars
        """
        return re.sub("\s+", " ", s)

    def build_file(self):
        document = Document()

        # Styles
        style = document.styles.add_style("Label Bold", WD_STYLE_TYPE.PARAGRAPH)
        style.font.bold = True
        style = document.styles.add_style("Table Cell", WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(9)
        style = document.styles.add_style(
            "Table Cell Bold", WD_STYLE_TYPE.PARAGRAPH
        )
        style.font.size = Pt(9)
        style.font.bold = True

        p = document.add_paragraph("Ref. Number")
        document.add_heading(self.context.getId(), 0)

        p = document.add_paragraph("")
        table = document.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Country"
        hdr_cells[0].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[1].text = "Sector"
        hdr_cells[1].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[2].text = "Gases"
        hdr_cells[2].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[3].text = "Fuel"
        hdr_cells[3].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[4].text = "Inventory year"
        hdr_cells[4].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[5].text = "Phase"
        hdr_cells[5].paragraphs[0].style = "Table Cell Bold"

        row_cells = table.add_row().cells
        row_cells[0].text = self.context.country_value() or ""
        row_cells[0].paragraphs[0].style = "Table Cell"
        row_cells[1].text = self.context.ghg_source_sectors_value() or ""
        row_cells[1].paragraphs[0].style = "Table Cell"
        row_cells[2].text = self.context.gas_value() or ""
        row_cells[2].paragraphs[0].style = "Table Cell"
        row_cells[3].text = self.context.fuel or ""
        row_cells[3].paragraphs[0].style = "Table Cell"
        row_cells[4].text = self.context.year or ""
        row_cells[4].paragraphs[0].style = "Table Cell"
        if self.context.observation_phase() == "phase1-observation":
            row_cells[5].text = "1"
            row_cells[5].paragraphs[0].style = "Table Cell"
        else:
            row_cells[5].text = "2"
            row_cells[5].paragraphs[0].style = "Table Cell"
        p = document.add_paragraph("")

        document.add_heading("Observation details", level=2)

        p = document.add_paragraph("")
        table = document.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Review Year"
        hdr_cells[1].text = "Parameter"
        hdr_cells[2].text = "Key category"
        hdr_cells[3].text = "Last update"
        hdr_cells[0].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[1].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[2].paragraphs[0].style = "Table Cell Bold"
        hdr_cells[3].paragraphs[0].style = "Table Cell Bold"

        row_cells = table.add_row().cells
        row_cells[0].text = "%s" % self.context.review_year or ""
        row_cells[1].text = self.context.parameter_value() or ""
        if self.context.ms_key_catagory:
            row_cells[2].text = "MS Key category"
        elif self.context.ms_key_catagory:
            row_cells[2].text = "EU Key category"
        else:
            row_cells[2].text = ""
        row_cells[3].text = self.context.modified().strftime(
            "%d %b %Y, %H:%M CET"
        )
        row_cells[0].paragraphs[0].style = "Table Cell"
        row_cells[1].paragraphs[0].style = "Table Cell"
        row_cells[2].paragraphs[0].style = "Table Cell"
        row_cells[3].paragraphs[0].style = "Table Cell"
        p = document.add_paragraph("")

        p = document.add_paragraph("Description flags", style="Label Bold")
        p = document.add_paragraph(self.context.highlight_value())
        p = document.add_paragraph(
            "Short description by expert/reviewer", style="Label Bold"
        )
        p = document.add_paragraph(self.context.text)

        if self.context.get_status() == "phase1-close-requested":
            document.add_heading("Finish observation", level=2)
            document.add_heading("Observation Finish Requested", level=3)
            p = document.add_paragraph(
                "SE comments on finish observation request:", style="Label Bold"
            )
            p = document.add_paragraph(self.context.closing_comments)
            if self.context.closing_deny_comments:
                document.add_heading(
                    "Observation Finalisation previously denied", level=3
                )
                p = document.add_paragraph(
                    "QE comments on finish observation request:",
                    style="Label Bold",
                )
                p = document.add_paragraph(self.context.closing_deny_comments)

        if self.context.get_status() == "phase2-close-requested":
            document.add_heading("Finish observation", level=2)
            document.add_heading("Observation Finish Requested", level=3)
            p = document.add_paragraph(
                "RE comments on finish observation request:", style="Label Bold"
            )
            p = document.add_paragraph(self.context.closing_comments_phase2)
            if self.context.closing_deny_comments_phase2:
                document.add_heading(
                    "Observation Finalisation previously denied", level=3
                )
                p = document.add_paragraph(
                    "LR comments on finish observation request:",
                    style="Label Bold",
                )
                p = document.add_paragraph(
                    self.context.closing_deny_comments_phase2
                )

        conclusion_2 = self.get_conclusion_phase2()
        if conclusion_2:
            document.add_page_break()
            document.add_heading("Conclusions Step 2", level=2)

            p = document.add_paragraph(
                "Final status of observation:", style="Label Bold"
            )
            p = document.add_paragraph(conclusion_2.reason_value())
            p = document.add_paragraph("Recommendation:", style="Label Bold")
            p = document.add_paragraph(conclusion_2.text)

            if self.context.closing_deny_comments_phase2:
                document.add_heading("Observation Finalisation denied", level=3)
                p = document.add_paragraph(
                    "LR comments on finish observation request:",
                    style="Label Bold",
                )
                p = document.add_paragraph(
                    self.context.closing_deny_comments_phase2
                )

        conclusion = self.get_conclusion()
        if conclusion:
            if conclusion_2 is not None and not conclusion_2():
                document.add_page_break()
            document.add_heading("Conclusions Step 1", level=2)

            p = document.add_paragraph(
                "Status of Observation:", style="Label Bold"
            )
            p = document.add_paragraph(conclusion.reason_value())
            p = document.add_paragraph("Internal Note:", style="Label Bold")
            p = document.add_paragraph(conclusion.text)

            if self.context.closing_deny_comments:
                document.add_heading("Observation Finalisation denied", level=3)
                p = document.add_paragraph(
                    "QE comments on finish observation request:",
                    style="Label Bold",
                )
                p = document.add_paragraph(self.context.closing_deny_comments)
            document.add_page_break()

        chats = self.get_chat()
        if chats:
            document.add_heading("Q&A", level=2)
            for chat in chats:
                date = chat.effective_date
                sent_info = "Sent on: %s"
                if not date:
                    date = chat.modified()
                    sent_info = "Updated on: %s"

                if chat.portal_type.lower() == "comment":
                    p = document.add_paragraph(
                        "> %s" % self.strip_special_chars(chat.text)
                    )
                    p = document.add_paragraph(
                        "From expert/reviewer To Member State \t\t %s"
                        % (sent_info % date.strftime("%d %b %Y, %H:%M CET"))
                    )

                if chat.portal_type.lower() == "commentanswer":
                    p = document.add_paragraph(
                        "< %s" % self.strip_special_chars(chat.text)
                    )
                    p = document.add_paragraph(
                        "From Member State To expert/reviewer \t\t %s"
                        % (sent_info % date.strftime("%d %b %Y, %H:%M CET"))
                    )

        return document

    def render(self):
        """ Export current filters observation in xls
        """
        document = self.build_file()

        response = self.request.response
        response.setHeader(
            "content-type",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response.setHeader(
            "Content-disposition",
            "attachment;filename=" + self.context.getId() + ".docx",
        )

        f = StringIO()
        document.save(f)
        f.seek(0)
        response.setHeader("Content-Length", len(f.getvalue()))
        response.write(f.getvalue())


class AddQuestionForm(Form):

    ignoreContext = True
    fields = field.Fields(IComment).select("text")

    @button.buttonAndHandler(u"Save question")
    def create_question(self, action):
        context = aq_inner(self.context)
        text = self.request.form.get("form.widgets.text", "")
        if not text.strip():
            raise ActionExecutionError(Invalid(u"Question text is empty"))

        qs = self.context.get_values_cat("Question")
        if qs:
            question = qs[0]
        else:
            q_id = context.invokeFactory(
                type_name="Question", id="question-1", title="Question 1",
            )
            question = context.get(q_id)

        id = str(int(time()))
        item_id = question.invokeFactory(type_name="Comment", id=id,)
        comment = question.get(item_id)
        comment.text = text

        return self.request.response.redirect(context.absolute_url())

    def updateWidgets(self):
        super(AddQuestionForm, self).updateWidgets()
        self.widgets["text"].rows = 15

    def updateActions(self):
        super(AddQuestionForm, self).updateActions()
        for k in self.actions.keys():
            self.actions[k].addClass("standardButton")
            self.actions[k].addClass("defaultWFButton")


class ModificationForm(dexterity.EditForm):
    grok.name("modifications")
    grok.context(IObservation)
    grok.require("cmf.ModifyPortalContent")

    def updateFields(self):
        super(ModificationForm, self).updateFields()

        user = api.user.get_current()
        roles = api.user.get_roles(username=user.getId(), obj=self.context)
        fields = []
        # XXX Needed? Edit rights are controlled by the WF
        if "ReviewerPhase1" in roles or "ReviewerPhase2" in roles:
            fields = [
                f
                for f in field.Fields(IObservation)
                if f
                not in [
                    "country",
                    "crf_code",
                    "review_year",
                    "technical_corrections",
                    "closing_comments",
                    "closing_deny_comments",
                    "closing_comments_phase2",
                    "closing_deny_comments_phase2",
                ]
            ]
        elif "QualityExpert" in roles or "LeadReviewer" in roles:
            fields = ["text", "highlight"]

        elif "Manager" in roles:
            fields = [f for f in field.Fields(IObservation)]

        self.fields = field.Fields(IObservation).select(*fields)
        self.groups = [
            g for g in self.groups if g.label == "label_schema_default"
        ]
        if "parameter" in fields:
            self.fields["parameter"].widgetFactory = CheckBoxFieldWidget
        if "highlight" in fields:
            self.fields["highlight"].widgetFactory = CheckBoxFieldWidget
        if "gas" in fields:
            self.fields["gas"].widgetFactory = CheckBoxFieldWidget

    def updateWidgets(self):
        super(ModificationForm, self).updateWidgets()
        if "review_year" in self.widgets and not Observation.is_secretariat():
            self.widgets["review_year"].readonly = "readonly"

    def updateActions(self):
        super(ModificationForm, self).updateActions()
        for k in self.actions.keys():
            self.actions[k].addClass("standardButton")


class EditForm(ModificationForm):
    grok.name("edit")


class AddAnswerForm(Form):

    ignoreContext = True
    fields = field.Fields(ICommentAnswer).select("text")

    @button.buttonAndHandler(u"Save answer")
    def add_answer(self, action):
        text = self.request.form.get("form.widgets.text", "")
        if not text.strip():
            raise ActionExecutionError(Invalid(u"Answer text is empty"))
        observation = aq_inner(self.context)
        questions = [
            q for q in observation.values() if q.portal_type == "Question"
        ]
        if questions:
            context = questions[0]
        else:
            raise ActionExecutionError(Invalid(u"Invalid context"))

        if context.has_answers():
            raise ActionExecutionError(
                Invalid(u"Question is already answered!")
            )

        id = str(int(time()))
        item_id = context.invokeFactory(type_name="CommentAnswer", id=id,)
        comment = context.get(item_id)
        comment.text = text
        if context.get_state_api().startswith("phase1-"):
            action = "phase1-add-answer"
        elif context.get_state_api().startswith("phase2-"):
            action = "phase2-add-answer"
        else:
            raise ActionExecutionError(Invalid(u"Invalid context"))

        api.content.transition(obj=context, transition=action)

        return self.request.response.redirect(observation.absolute_url())

    def updateWidgets(self):
        super(AddAnswerForm, self).updateWidgets()
        self.widgets["text"].rows = 15

    def updateActions(self):
        super(AddAnswerForm, self).updateActions()
        for k in self.actions.keys():
            self.actions[k].addClass("standardButton")


class AddAnswerAndRequestComments(grok.View):
    grok.context(IObservation)
    grok.name("add-answer-and-request-comments")
    grok.require("zope2.View")

    def render(self):
        observation = aq_inner(self.context)
        questions = [
            q for q in observation.values() if q.portal_type == "Question"
        ]
        if questions:
            context = questions[0]
        else:
            raise ActionExecutionError(Invalid(u"Invalid context"))

        comments = [q for q in context.values() if q.portal_type == "Comment"]
        answers = [
            q for q in context.values() if q.portal_type == "CommentAnswer"
        ]

        if len(comments) == len(answers):
            status = IStatusMessage(self.request)
            msg = _(u"There is a draft answer created for the question.")
            status.addStatusMessage(msg, "error")
            return self.request.response.redirect(observation.absolute_url())

        context = questions[0]

        text = u"For MS coordinator: please draft, edit and finalise your consolidated reply here."

        id = str(int(time()))
        item_id = context.invokeFactory(type_name="CommentAnswer", id=id,)
        comment = context.get(item_id)
        comment.text = text
        if context.get_state_api().startswith("phase1-"):
            action = "phase1-assign-answerer"
        elif api.content.get_state(context).startswith("phase2-"):
            action = "phase2-assign-answerer"
        else:
            raise ActionExecutionError(Invalid(u"Invalid context"))
        url = "%s/assign_answerer_form?workflow_action=%s&comment=%s" % (
            context.absolute_url(),
            action,
            item_id,
        )
        return self.request.response.redirect(url)


class AddCommentForm(Form):

    ignoreContext = True
    fields = field.Fields(IComment).select("text")

    @button.buttonAndHandler(u"Add question")
    def create_question(self, action):
        observation = aq_inner(self.context)
        questions = [
            q for q in observation.values() if q.portal_type == "Question"
        ]
        if questions:
            context = questions[0]
        else:
            raise

        text = self.request.form.get("form.widgets.text", "")
        create_comment(text, context)

        return self.request.response.redirect(observation.absolute_url())

    def updateWidgets(self):
        super(AddCommentForm, self).updateWidgets()
        self.widgets["text"].rows = 15

    def updateActions(self):
        super(AddCommentForm, self).updateActions()
        for k in self.actions.keys():
            self.actions[k].addClass("standardButton")


class AddConclusionForm(Form):
    ignoreContext = True
    fields = field.Fields(IConclusion).select("text", "closing_reason")

    @button.buttonAndHandler(u"Add conclusion")
    def create_conclusion(self, action):
        context = aq_inner(self.context)
        id = str(int(time()))
        item_id = context.invokeFactory(type_name="Conclusion", id=id,)
        text = self.request.form.get("form.widgets.text", "")
        comment = context.get(item_id)
        comment.text = text
        reason = self.request.form.get("form.widgets.closing_reason")
        comment.closing_reason = reason[0]
        adapted = IAllowDiscussion(comment)
        adapted.allow_discussion = True

        return self.request.response.redirect(context.absolute_url())

    def updateWidgets(self):
        super(AddConclusionForm, self).updateWidgets()
        self.widgets["text"].rows = 15

    def updateActions(self):
        super(AddConclusionForm, self).updateActions()
        for k in self.actions.keys():
            self.actions[k].addClass("standardButton")


class EditConclusionAndCloseComments(grok.View):
    grok.name("edit-conclusions-and-close-comments")
    grok.context(IObservation)
    grok.require("zope2.View")

    def update(self):
        # Some checks:
        waction = self.request.get("workflow_action")
        if waction != "phase1-finish-comments":
            status = IStatusMessage(self.request)
            msg = u"There was an error, try again please"
            status.addStatusMessage(msg, "error")

    def render(self):
        # Execute the transition
        api.content.transition(
            obj=self.context, transition="phase1-finish-comments"
        )
        conclusions = self.context.get_values_cat("Conclusion")
        conclusion = conclusions[0]
        url = "%s/edit" % conclusion.absolute_url()
        return self.request.response.redirect(url)


class EditConclusionP2AndCloseComments(grok.View):
    grok.name("edit-conclusions-and-close-comments-phase2")
    grok.context(IObservation)
    grok.require("zope2.View")

    def update(self):
        # Some checks:
        waction = self.request.get("workflow_action")
        if waction != "phase2-finish-comments":
            status = IStatusMessage(self.request)
            msg = u"There was an error, try again please"
            status.addStatusMessage(msg, "error")

    def render(self):
        # Execute the transition
        api.content.transition(
            obj=self.context, transition="phase2-finish-comments"
        )
        conclusions = self.context.get_values_cat("ConclusionsPhase2")
        conclusion = conclusions[0]
        url = "%s/edit" % conclusion.absolute_url()
        return self.request.response.redirect(url)


class EditHighlightsForm(dexterity.EditForm):
    grok.name("edit-highlights")
    grok.context(IObservation)
    grok.require("cmf.ModifyPortalContent")

    def updateFields(self):
        super(EditHighlightsForm, self).updateFields()
        self.fields = field.Fields(IObservation).select("highlight")
        self.fields["highlight"].widgetFactory = CheckBoxFieldWidget
        self.groups = [
            g for g in self.groups if g.label == "label_schema_default"
        ]


class AddConclusions(grok.View):
    grok.context(IObservation)
    grok.name("add-conclusions")
    grok.require("zope2.View")

    def render(self):
        context = aq_inner(self.context)
        workflow = getToolByName(context, "portal_workflow")
        transitions_available = [
            action["id"] for action in workflow.listActions(object=context)
        ]
        if context.get_status().startswith("phase1-"):
            current_user_id = api.user.get_current().getId()
            user_roles = api.user.get_roles(
                username=current_user_id, obj=context
            )
            if "ReviewerPhase1" in user_roles:
                cs = self.context.get_values_cat("Conclusion")
                if cs:
                    conclusion = cs[0]
                    if "ReviewerPhase1" in user_roles:
                        url = conclusion.absolute_url() + "/edit"
                else:
                    # with api.env.adopt_roles(['ReviewerPhase1']):
                    id = context.invokeFactory(
                        id=str(int(time())), type_name="Conclusion", text=u""
                    )
                    cs = self.context.get_values_cat("Conclusion")
                    conclusion = cs[0]
                    url = conclusion.absolute_url() + "/edit"
                    # url = '%s/++add++Conclusion' % context.absolute_url()
            else:
                cs = self.context.get_values_cat("Conclusion")
                if cs:
                    conclusion = cs[0]
                    url = conclusion.absolute_url() + "/edit"
                else:
                    # with api.env.adopt_roles(['ReviewerPhase1']):
                    id = context.invokeFactory(
                        id=str(int(time())), type_name="Conclusion", text=u""
                    )
                    cs = self.context.get_values_cat("Conclusion")
                    conclusion = cs[0]
                    url = conclusion.absolute_url() + "/edit"
            if "phase1-draft-conclusions" in transitions_available:
                api.content.transition(
                    obj=context, transition="phase1-draft-conclusions"
                )

        elif context.get_status().startswith("phase2-"):
            current_user_id = api.user.get_current().getId()
            user_roles = api.user.get_roles(
                username=current_user_id, obj=context
            )
            if "ReviewerPhase2" in user_roles:
                csp2 = self.context.get_values_cat("ConclusionsPhase2")
                if csp2:
                    conclusionsphase2 = csp2[0]
                    url = conclusionsphase2.absolute_url() + "/edit"

                else:
                    # with api.env.adopt_roles(['ReviewerPhase2']):
                    id = context.invokeFactory(
                        id=str(int(time())),
                        type_name="ConclusionsPhase2",
                        text=u"",
                    )
                    cs = self.context.get_values_cat("ConclusionsPhase2")
                    conclusion = cs[0]
                    url = conclusion.absolute_url() + "/edit"
                    # url = '%s/++add++ConclusionsPhase2' % context.absolute_url()

            else:
                csp2 = self.context.get_values_cat("ConclusionsPhase2")
                if csp2:
                    conclusionsphase2 = csp2[0]
                    url = conclusionsphase2.absolute_url() + "/edit"

                else:
                    # with api.env.adopt_roles(['ReviewerPhase2']):
                    id = context.invokeFactory(
                        id=str(int(time())),
                        type_name="ConclusionsPhase2",
                        text=u"",
                    )
                    cs = self.context.get_values_cat("ConclusionsPhase2")
                    conclusion = cs[0]
                    url = conclusion.absolute_url() + "/edit"

            if "phase2-draft-conclusions" in transitions_available:
                api.content.transition(
                    obj=context, transition="phase2-draft-conclusions"
                )
        else:
            raise ActionExecutionError(Invalid(u"Invalid context"))

        self.context.reindexObject()
        return self.request.response.redirect(url)


def create_comment(text, question):
    id = str(int(time()))
    item_id = question.invokeFactory(type_name="Comment", id=id)
    comment = question.get(item_id)
    comment.text = text
    return comment
